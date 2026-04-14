#!/usr/bin/env python3
"""
MD to DOCX Converter by Jair Lima
Converts Markdown files to Word DOCX with proper formatting.

Usage:
  md2docx                    # Convert all .md in current folder
  md2docx arquivo.md         # Convert specific file
  md2docx --force arquivo.md # Overwrite even if DOCX already exists
"""

import sys
import os
import re
import argparse
import threading
import itertools
import time
import tempfile
import urllib.request
from pathlib import Path

VERSION = "3.3"

# Fix Windows terminal encoding
if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

import mistune
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT


# ─────────────────────────────────────────────────────────────────────────────
# Spinner (feedback visual durante conversão)
# ─────────────────────────────────────────────────────────────────────────────

class Spinner:
    """Context manager que exibe um spinner animado na linha atual."""

    FRAMES = ["|", "/", "-", "\\"]

    def __init__(self, msg: str = "Convertendo"):
        self.msg = msg
        self._stop = threading.Event()
        self._thread = threading.Thread(target=self._spin, daemon=True)

    def _spin(self):
        for frame in itertools.cycle(self.FRAMES):
            if self._stop.is_set():
                break
            sys.stdout.write(f"\r  [{frame}]  {self.msg}...")
            sys.stdout.flush()
            time.sleep(0.1)

    def __enter__(self):
        self._thread.start()
        return self

    def __exit__(self, *args):
        self._stop.set()
        self._thread.join()
        sys.stdout.write("\r" + " " * (len(self.msg) + 14) + "\r")
        sys.stdout.flush()


# ─────────────────────────────────────────────────────────────────────────────
# Document styles setup
# ─────────────────────────────────────────────────────────────────────────────

def setup_styles(doc: Document):
    """Configure all document styles for proper DOCX output."""

    # ── Normal (base) ────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # ── Headings ─────────────────────────────────────────────────────────────
    heading_config = [
        ("Heading 1", 22, True, RGBColor(0x1F, 0x39, 0x64), Pt(12), Pt(6)),
        ("Heading 2", 16, True, RGBColor(0x2E, 0x74, 0xB5), Pt(10), Pt(4)),
        ("Heading 3", 13, True, RGBColor(0x1F, 0x39, 0x64), Pt(8),  Pt(4)),
        ("Heading 4", 12, True, RGBColor(0x2E, 0x74, 0xB5), Pt(6),  Pt(2)),
        ("Heading 5", 11, True, RGBColor(0x40, 0x40, 0x40), Pt(4),  Pt(2)),
        ("Heading 6", 11, False, RGBColor(0x70, 0x70, 0x70), Pt(4), Pt(2)),
    ]
    for name, size, bold, color, space_before, space_after in heading_config:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = space_before
        style.paragraph_format.space_after = space_after
        style.paragraph_format.keep_with_next = True

    # ── Code (inline/block character style) ──────────────────────────────────
    if "Code Char" not in [s.name for s in doc.styles]:
        code_char = doc.styles.add_style("Code Char", 2)  # 2 = character style
    else:
        code_char = doc.styles["Code Char"]
    code_char.font.name = "Courier New"
    code_char.font.size = Pt(10)

    # ── Code Block (paragraph) ────────────────────────────────────────────────
    if "Code Block" not in [s.name for s in doc.styles]:
        code_block = doc.styles.add_style("Code Block", 1)
        code_block.base_style = doc.styles["Normal"]
    else:
        code_block = doc.styles["Code Block"]
    code_block.font.name = "Courier New"
    code_block.font.size = Pt(9.5)
    code_block.paragraph_format.space_before = Pt(6)
    code_block.paragraph_format.space_after = Pt(6)
    code_block.paragraph_format.left_indent = Cm(0.5)
    code_block.paragraph_format.right_indent = Cm(0.5)

    # ── Block Quote ───────────────────────────────────────────────────────────
    if "Block Quote" not in [s.name for s in doc.styles]:
        bq = doc.styles.add_style("Block Quote", 1)
        bq.base_style = doc.styles["Normal"]
    else:
        bq = doc.styles["Block Quote"]
    bq.font.name = "Calibri"
    bq.font.size = Pt(11)
    bq.font.italic = True
    bq.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    bq.paragraph_format.left_indent = Cm(1.0)
    bq.paragraph_format.right_indent = Cm(1.0)
    bq.paragraph_format.space_before = Pt(4)
    bq.paragraph_format.space_after = Pt(4)


def set_paragraph_shading(para, fill_hex: str):
    """Add background shading to a paragraph."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_left_border(para, color_hex: str = "2E74B5", size: int = 24):
    """Add a left border to a paragraph (blockquote style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def add_hyperlink(para, text: str, url: str):
    """Add a hyperlink run to a paragraph."""
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    hyperlink.append(r)
    para._p.append(hyperlink)
    return hyperlink


def add_horizontal_rule(doc: Document):
    """Add a horizontal rule paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def set_table_style(table):
    """Apply clean table borders and shading."""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # Table borders
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def shade_table_header(row):
    """Apply header shading to a table row."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "DEEAF1")
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


# ─────────────────────────────────────────────────────────────────────────────
# Table column auto-fit
# ─────────────────────────────────────────────────────────────────────────────

# Usable page width: A4 (21cm) minus 2 × 2.54cm margins = 15.92cm
_USABLE_WIDTH_TWIPS = int(Cm(15.92) / 635)  # 1 twip = 635 EMU


def _auto_fit_table_columns(table, all_row_tokens, text_extractor):
    """
    Distribute column widths proportionally based on the longest text found
    in each column. Ensures no column is narrower than MIN_CHARS equivalent.
    """
    MIN_CHARS = 6
    num_cols = len(table.columns)
    col_max = [MIN_CHARS] * num_cols

    for row_tok in all_row_tokens:
        if not isinstance(row_tok, dict):
            continue
        for col_idx, cell_tok in enumerate(row_tok.get("children", [])):
            if col_idx >= num_cols:
                break
            if isinstance(cell_tok, dict):
                text = text_extractor(cell_tok.get("children", []))
                col_max[col_idx] = max(col_max[col_idx], len(text))

    total = sum(col_max)
    widths = [max(1, int(_USABLE_WIDTH_TWIPS * w / total)) for w in col_max]

    # Fix rounding drift so widths sum exactly to usable width
    diff = _USABLE_WIDTH_TWIPS - sum(widths)
    widths[0] += diff

    # Apply to tblGrid
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(2, tblGrid)
    for gc in list(tblGrid):
        tblGrid.remove(gc)
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)

    # Apply to each cell
    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            if col_idx >= len(widths):
                break
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(widths[col_idx]))
            tcW.set(qn("w:type"), "dxa")


# ─────────────────────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────────────────────

def extract_cover(md_text: str):
    """
    Detect a cover block at the start of the MD:
      # Title
      ## Subtitle          (optional)
      *Author line*        (optional, one or more)
      ---                  (optional separator)

    Returns (cover_dict | None, remaining_md_text).
    """
    m = re.match(
        r'^# ([^\n]+)\n'              # H1 title (required)
        r'(?:## ([^\n]+)\n)?'         # H2 subtitle (optional)
        r'(?:[ \t]*\n)*'              # blank lines
        r'((?:\*[^\n]+\*[ \t]*\n)*)'  # italic lines (optional)
        r'(?:[ \t]*\n)*'              # blank lines
        r'(?:---+[ \t]*\n)?',         # separator (optional)
        md_text,
    )
    if not m or not m.group(1).strip():
        return None, md_text

    title    = m.group(1).strip()
    subtitle = m.group(2).strip() if m.group(2) else None
    meta_raw = m.group(3) or ""
    meta_lines = [l.strip().strip("*").strip()
                  for l in meta_raw.splitlines() if l.strip()]

    cover = {"title": title, "subtitle": subtitle, "meta_lines": meta_lines}
    return cover, md_text[m.end():]


def add_cover_page(doc: Document, cover: dict):
    """Render a centered cover page and add a page break."""
    # Vertical padding (push content towards vertical center)
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # ── Book title ────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(10)
    run = title_para.add_run(cover["title"])
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    run.font.name = "Calibri"

    # ── Subtitle ──────────────────────────────────────────────────────────────
    if cover.get("subtitle"):
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_after = Pt(40)
        run = sub_para.add_run(cover["subtitle"])
        run.italic = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.name = "Calibri"

    # ── Separator line ────────────────────────────────────────────────────────
    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(20)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Author / meta lines ───────────────────────────────────────────────────
    for line in cover.get("meta_lines", []):
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.paragraph_format.space_after = Pt(4)
        run = meta_para.add_run(line)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run.font.name = "Calibri"

    # ── Page break ────────────────────────────────────────────────────────────
    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Table of Contents
# ─────────────────────────────────────────────────────────────────────────────

def add_toc(doc: Document):
    """Insert a Word TOC field (requires Ctrl+A, F9 in Word to populate)."""
    title = doc.add_paragraph("Sumário", style="Heading 1")
    title.paragraph_format.page_break_before = False

    para = doc.add_paragraph()
    run = para.add_run()

    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "(Atualize o campo: Ctrl+A → F9)"

    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(placeholder)
    run._r.append(fldChar_end)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Footer with page numbers
# ─────────────────────────────────────────────────────────────────────────────

def add_footer_page_numbers(doc: Document):
    """Add centred page number to footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.clear()

        run = para.add_run()
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        instrText.text = " PAGE "

        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)


# ─────────────────────────────────────────────────────────────────────────────
# Inline formatting parser
# ─────────────────────────────────────────────────────────────────────────────

def apply_inline(para, text: str, base_bold=False, base_italic=False,
                 base_size=None, base_color=None, font_name=None):
    """
    Parse inline markdown (bold, italic, strikethrough, inline code, links)
    and add runs to the paragraph.
    """
    # Pattern order matters — more specific first
    pattern = re.compile(
        r"(\*\*\*|___)"           # bold+italic (***  or ___)
        r"|(\*\*|__)"             # bold
        r"|(\*|_)"                # italic
        r"|(~~)"                  # strikethrough
        r"|(`.+?`)"               # inline code
        r"|(\[([^\]]*)\]\(([^)]*)\))"  # link [text](url)
        r"|(<([^>]+)>)"           # autolink <url>
        r"|(\\\S)"                # escaped char
    )

    i = 0
    segments = []
    for m in pattern.finditer(text):
        if m.start() > i:
            segments.append(("text", text[i:m.start()]))
        if m.group(1):   # bold+italic
            segments.append(("toggle_bi", m.group(1)))
        elif m.group(2): # bold
            segments.append(("toggle_b", m.group(2)))
        elif m.group(3): # italic
            segments.append(("toggle_i", m.group(3)))
        elif m.group(4): # strikethrough
            segments.append(("toggle_s", m.group(4)))
        elif m.group(5): # inline code
            code_text = m.group(5)[1:-1]  # strip backticks
            segments.append(("code", code_text))
        elif m.group(6): # link
            segments.append(("link", m.group(7), m.group(8)))
        elif m.group(9): # autolink
            url = m.group(10)
            segments.append(("link", url, url))
        elif m.group(11): # escaped
            segments.append(("text", m.group(11)[1:]))
        i = m.end()

    if i < len(text):
        segments.append(("text", text[i:]))

    # Render segments with toggle state
    b, it, s = base_bold, base_italic, False
    for seg in segments:
        kind = seg[0]
        if kind == "text":
            if not seg[1]:
                continue
            run = para.add_run(seg[1])
            run.bold = b
            run.italic = it
            run.font.strike = s
            if base_size:
                run.font.size = base_size
            if base_color:
                run.font.color.rgb = base_color
            if font_name:
                run.font.name = font_name
        elif kind == "toggle_bi":
            b = not b
            it = not it
        elif kind == "toggle_b":
            b = not b
        elif kind == "toggle_i":
            it = not it
        elif kind == "toggle_s":
            s = not s
        elif kind == "code":
            run = para.add_run(seg[1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            # light gray highlight for inline code
            rPr = run._r.get_or_add_rPr()
            highlight = OxmlElement("w:highlight")
            highlight.set(qn("w:val"), "lightGray")
            rPr.append(highlight)
        elif kind == "link":
            link_text, url = seg[1], seg[2]
            try:
                add_hyperlink(para, link_text or url, url)
            except Exception:
                run = para.add_run(link_text or url)
                run.font.color.rgb = RGBColor(0x00, 0x56, 0xB2)


# ─────────────────────────────────────────────────────────────────────────────
# AST → DOCX renderer
# ─────────────────────────────────────────────────────────────────────────────

class DocxRenderer(mistune.BaseRenderer):
    """Renders a mistune AST directly into a python-docx Document."""

    def __init__(self, doc: Document, md_path: Path):
        super().__init__()
        self.doc = doc
        self.md_path = md_path  # for resolving relative image paths
        self._list_level = 0
        self._list_type_stack = []  # "bullet" or "number"
        self._in_table_header = False
        self._h1_count = 0          # tracks H1s for page-break logic
        self._footnotes: list[tuple[int, list]] = []  # (id, inline_tokens)

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _inline_tokens_to_text(self, tokens) -> str:
        """Flatten inline tokens to plain text (for tables / headings fallback)."""
        parts = []
        for tok in tokens:
            if isinstance(tok, dict):
                t = tok.get("type", "")
                if t == "text":
                    parts.append(tok.get("raw", ""))
                elif t in ("strong", "em", "codespan", "del"):
                    parts.append(self._inline_tokens_to_text(tok.get("children", [])))
                elif t == "softline":
                    parts.append(" ")
                elif t == "linebreak":
                    parts.append("\n")
                elif t == "link":
                    parts.append(self._inline_tokens_to_text(tok.get("children", [])))
                elif t == "image":
                    parts.append(tok.get("attrs", {}).get("alt", ""))
                elif t == "raw":
                    raw = tok.get("raw", "")
                    parts.append(re.sub(r"<[^>]+>", "", raw))
            elif isinstance(tok, str):
                parts.append(tok)
        return "".join(parts)

    def _render_inline_to_para(self, para, tokens, **kwargs):
        """Render a list of inline tokens to runs in a paragraph."""
        for tok in tokens:
            if isinstance(tok, dict):
                t = tok.get("type", "")
                if t == "text":
                    apply_inline(para, tok.get("raw", ""), **kwargs)
                elif t == "softline":
                    para.add_run(" ")
                elif t == "linebreak":
                    para.add_run().add_break()
                elif t == "strong":
                    self._render_inline_to_para(para, tok.get("children", []),
                                                 base_bold=True, **{k: v for k, v in kwargs.items() if k != "base_bold"})
                elif t == "em":
                    self._render_inline_to_para(para, tok.get("children", []),
                                                 base_italic=True, **{k: v for k, v in kwargs.items() if k != "base_italic"})
                elif t == "del":
                    child_tokens = tok.get("children", [])
                    for ct in child_tokens:
                        if isinstance(ct, dict) and ct.get("type") == "text":
                            run = para.add_run(ct.get("raw", ""))
                            run.font.strike = True
                elif t == "codespan":
                    run = para.add_run(tok.get("raw", ""))
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                    rPr = run._r.get_or_add_rPr()
                    highlight = OxmlElement("w:highlight")
                    highlight.set(qn("w:val"), "lightGray")
                    rPr.append(highlight)
                elif t == "link":
                    url = tok.get("attrs", {}).get("url", "#")
                    link_text = self._inline_tokens_to_text(tok.get("children", []))
                    try:
                        add_hyperlink(para, link_text or url, url)
                    except Exception:
                        run = para.add_run(link_text or url)
                        run.font.color.rgb = RGBColor(0x00, 0x56, 0xB2)
                elif t == "image":
                    attrs = tok.get("attrs", {})
                    src = attrs.get("url", "")
                    alt = attrs.get("alt", "")
                    self._add_image_to_para(para, src, alt)
                elif t == "footnote_ref":
                    idx = tok.get("attrs", {}).get("index", 1)
                    r = OxmlElement("w:r")
                    rPr = OxmlElement("w:rPr")
                    rStyle = OxmlElement("w:rStyle")
                    rStyle.set(qn("w:val"), "FootnoteReference")
                    rPr.append(rStyle)
                    r.append(rPr)
                    fn_ref = OxmlElement("w:footnoteReference")
                    fn_ref.set(qn("w:id"), str(idx))
                    r.append(fn_ref)
                    para._p.append(r)
                elif t == "raw":
                    raw = tok.get("raw", "")
                    clean = re.sub(r"<[^>]+>", "", raw)
                    if clean.strip():
                        para.add_run(clean)
            elif isinstance(tok, str):
                apply_inline(para, tok, **kwargs)

    def _add_image_to_para(self, para, src: str, alt: str = ""):
        """Try to add an image (local or remote URL); fall back to alt text."""
        # ── Remote URL ────────────────────────────────────────────────────────
        if src.startswith(("http://", "https://")):
            suffix = Path(src.split("?")[0]).suffix or ".png"
            tmp_path = None
            try:
                fd, tmp_name = tempfile.mkstemp(suffix=suffix)
                os.close(fd)
                tmp_path = Path(tmp_name)
                urllib.request.urlretrieve(src, str(tmp_path))
                run = para.add_run()
                run.add_picture(str(tmp_path), width=Inches(5.5))
                return
            except Exception:
                pass
            finally:
                if tmp_path and tmp_path.exists():
                    try:
                        tmp_path.unlink()
                    except Exception:
                        pass

        # ── Local file ────────────────────────────────────────────────────────
        else:
            img_path = Path(src)
            if not img_path.is_absolute():
                img_path = self.md_path.parent / src
            if img_path.exists():
                try:
                    run = para.add_run()
                    run.add_picture(str(img_path), width=Inches(5.5))
                    return
                except Exception:
                    pass

        display = f"[Imagem: {alt or src}]"
        run = para.add_run(display)
        run.italic = True
        run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

    # ── Block elements ────────────────────────────────────────────────────────

    def heading(self, token, state):
        level = token.get("attrs", {}).get("level", 1)
        children = token.get("children", [])
        style_name = f"Heading {min(level, 6)}"
        para = self.doc.add_paragraph(style=style_name)

        # Page break before every H1 except the very first
        if level == 1:
            if self._h1_count > 0:
                para.paragraph_format.page_break_before = True
            self._h1_count += 1

        self._render_inline_to_para(para, children)
        return ""

    def paragraph(self, token, state):
        children = token.get("children", [])
        para = self.doc.add_paragraph(style="Normal")
        self._render_inline_to_para(para, children)
        return ""

    def blank_line(self, token, state):
        return ""

    def thematic_break(self, token, state):
        add_horizontal_rule(self.doc)
        return ""

    def block_code(self, token, state):
        """Fenced or indented code block."""
        raw = token.get("raw", "")
        lines = raw.split("\n")
        if lines and lines[-1] == "":
            lines = lines[:-1]

        for i, line in enumerate(lines):
            para = self.doc.add_paragraph(style="Code Block")
            set_paragraph_shading(para, "F2F2F2")
            para.paragraph_format.space_before = Pt(0) if i > 0 else Pt(6)
            para.paragraph_format.space_after = Pt(0) if i < len(lines) - 1 else Pt(6)
            run = para.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        return ""

    def block_quote(self, token, state):
        """
        Render a blockquote.
        Detects Bible-verse pattern: text followed by softline + "— Reference"
        and renders with amber border + right-aligned reference line.
        """
        children = token.get("children", [])
        for child in children:
            if not isinstance(child, dict):
                continue
            ct = child.get("type", "")
            if ct != "paragraph":
                self._render_token(child, state)
                continue

            inline = child.get("children", [])

            # ── Detect citation pattern ───────────────────────────────────────
            # Look for a softline whose next text token starts with — (em dash)
            split_idx = None
            for i, tok in enumerate(inline):
                if not isinstance(tok, dict) or tok.get("type") != "softline":
                    continue
                for j in range(i + 1, len(inline)):
                    nxt = inline[j]
                    if not isinstance(nxt, dict):
                        break
                    nt = nxt.get("type", "")
                    if nt == "softline":
                        continue
                    if nt == "text":
                        raw = nxt.get("raw", "").strip()
                        if raw.startswith("\u2014") or raw.startswith("—"):
                            split_idx = i
                    break
                if split_idx is not None:
                    break

            if split_idx is not None:
                # ── Quote text (amber left border, slightly smaller) ───────────
                quote_para = self.doc.add_paragraph(style="Block Quote")
                set_left_border(quote_para, "8B6914", size=20)
                quote_para.paragraph_format.space_after = Pt(2)
                self._render_inline_to_para(
                    quote_para, inline[:split_idx],
                    base_size=Pt(10.5),
                    base_color=RGBColor(0x25, 0x25, 0x25),
                )

                # ── Reference line (right-aligned, gray, smaller) ─────────────
                ref_tokens = inline[split_idx + 1:]  # skip the softline itself
                ref_para = self.doc.add_paragraph(style="Block Quote")
                ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                ref_para.paragraph_format.space_before = Pt(0)
                ref_para.paragraph_format.space_after = Pt(10)
                self._render_inline_to_para(
                    ref_para, ref_tokens,
                    base_size=Pt(10),
                    base_italic=False,
                    base_color=RGBColor(0x60, 0x60, 0x60),
                )
            else:
                # ── Regular blockquote ────────────────────────────────────────
                para = self.doc.add_paragraph(style="Block Quote")
                set_left_border(para)
                self._render_inline_to_para(para, inline)

        return ""

    def list(self, token, state):
        ordered = token.get("attrs", {}).get("ordered", False)
        self._list_type_stack.append("number" if ordered else "bullet")
        self._list_level += 1
        children = token.get("children", [])
        for child in children:
            self._render_token(child, state)
        self._list_level -= 1
        self._list_type_stack.pop()
        return ""

    def list_item(self, token, state):
        list_type = self._list_type_stack[-1] if self._list_type_stack else "bullet"
        level = self._list_level  # 1-based

        # Task list support: checked = True/False, None = regular item
        checked = token.get("attrs", {}).get("checked", None)

        if list_type == "bullet":
            style = "List Bullet" if level == 1 else f"List Bullet {min(level, 3)}"
        else:
            style = "List Number" if level == 1 else f"List Number {min(level, 3)}"

        children = token.get("children", [])
        para = None
        checkbox_added = False

        def _add_checkbox(p):
            nonlocal checkbox_added
            if checked is not None and not checkbox_added:
                p.add_run("☑ " if checked else "☐ ")
                checkbox_added = True

        for child in children:
            if isinstance(child, dict):
                ct = child.get("type", "")
                if ct == "block_text":
                    para = self.doc.add_paragraph(style=style)
                    _add_checkbox(para)
                    self._render_inline_to_para(para, child.get("children", []))
                elif ct == "paragraph":
                    if para is None:
                        para = self.doc.add_paragraph(style=style)
                        _add_checkbox(para)
                        self._render_inline_to_para(para, child.get("children", []))
                    else:
                        p2 = self.doc.add_paragraph(style="Normal")
                        p2.paragraph_format.left_indent = Cm(level * 0.63)
                        self._render_inline_to_para(p2, child.get("children", []))
                elif ct == "list":
                    self._render_token(child, state)
                else:
                    self._render_token(child, state)
            elif isinstance(child, str) and child.strip():
                if para is None:
                    para = self.doc.add_paragraph(style=style)
                    _add_checkbox(para)
                    apply_inline(para, child)
        return ""

    def task_list_item(self, token, state):
        """Delegate to list_item — checked attr is already in token.attrs."""
        return self.list_item(token, state)

    def footnotes(self, token, state):
        """Collect footnote definitions (as inline token lists) for later XML attachment."""
        for child in token.get("children", []):
            if not isinstance(child, dict) or child.get("type") != "footnote_item":
                continue
            idx = child.get("attrs", {}).get("index", len(self._footnotes) + 1)
            inline_tokens = self._collect_footnote_tokens(child)
            self._footnotes.append((idx, inline_tokens))
        return ""

    def _collect_footnote_tokens(self, token) -> list:
        """Collect all inline tokens from a footnote_item, joining paragraphs with a space."""
        result = []
        for child in token.get("children", []):
            if not isinstance(child, dict):
                continue
            ct = child.get("type", "")
            if ct in ("paragraph", "block_text"):
                if result:
                    result.append({"type": "softline"})
                result.extend(child.get("children", []))
        return result

    def table(self, token, state):
        children = token.get("children", [])
        if not children:
            return ""

        head_rows = []
        body_rows = []
        for child in children:
            if isinstance(child, dict):
                if child.get("type") == "table_head":
                    for row in child.get("children", []):
                        head_rows.append(row)
                elif child.get("type") == "table_body":
                    for row in child.get("children", []):
                        body_rows.append(row)

        if not head_rows and not body_rows:
            return ""

        num_cols = max(
            len(r.get("children", [])) for r in (head_rows + body_rows)
            if isinstance(r, dict)
        ) if (head_rows or body_rows) else 1

        num_rows = len(head_rows) + len(body_rows)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_style(table)

        row_idx = 0
        for row_tok in head_rows:
            if isinstance(row_tok, dict):
                self._fill_table_row(table.rows[row_idx], row_tok, is_header=True)
                row_idx += 1

        for row_tok in body_rows:
            if isinstance(row_tok, dict):
                self._fill_table_row(table.rows[row_idx], row_tok, is_header=False)
                row_idx += 1

        if head_rows:
            shade_table_header(table.rows[0])

        _auto_fit_table_columns(
            table,
            head_rows + body_rows,
            self._inline_tokens_to_text,
        )

        self.doc.add_paragraph(style="Normal").paragraph_format.space_after = Pt(4)
        return ""

    def _fill_table_row(self, row, row_token, is_header: bool):
        cells = row_token.get("children", [])
        for col_idx, cell_tok in enumerate(cells):
            if col_idx >= len(row.cells):
                break
            cell = row.cells[col_idx]
            for p in cell.paragraphs:
                p._element.getparent().remove(p._element)

            para = cell.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)

            if isinstance(cell_tok, dict):
                align = cell_tok.get("attrs", {}).get("align")
                if align == "center":
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                inline_children = cell_tok.get("children", [])
                self._render_inline_to_para(para, inline_children,
                                             base_bold=is_header)

    def _render_token(self, token, state):
        """Dispatch a single token to its renderer method."""
        if not isinstance(token, dict):
            return
        t = token.get("type", "")
        method = getattr(self, t, None)
        if method:
            method(token, state)
        elif t in ("block_text",):
            para = self.doc.add_paragraph(style="Normal")
            self._render_inline_to_para(para, token.get("children", []))

    # ── Raw HTML (strip tags) ─────────────────────────────────────────────────
    def block_html(self, token, state):
        raw = token.get("raw", "")
        if re.search(r"<hr\s*/?>", raw, re.I):
            add_horizontal_rule(self.doc)
            return ""
        clean = re.sub(r"<[^>]+>", "", raw).strip()
        if clean:
            para = self.doc.add_paragraph(style="Normal")
            para.add_run(clean)
        return ""

    def inline_html(self, token, state):
        return ""

    # ── Finalize ──────────────────────────────────────────────────────────────
    def finalize(self, data, state):
        return ""


# ─────────────────────────────────────────────────────────────────────────────
# DOCX → MD helpers
# ─────────────────────────────────────────────────────────────────────────────

_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _w(tag: str) -> str:
    return f"{{{_NS_W}}}{tag}"


def _attr_is_on(el, tag: str) -> bool:
    """Return True if <w:tag> exists inside el and is NOT explicitly set to 0/false."""
    child = el.find(_w(tag))
    if child is None:
        return False
    val = child.get(_w("val"), "true")
    return val.lower() not in ("0", "false")


def _run_el_to_md(r_el) -> str:
    """Convert a <w:r> XML element to inline Markdown text."""
    rpr = r_el.find(_w("rPr"))
    texts = r_el.findall(_w("t"))
    text = "".join(t.text or "" for t in texts)
    if not text:
        return ""

    bold = italic = strike = is_code = False

    if rpr is not None:
        bold   = _attr_is_on(rpr, "b")
        italic = _attr_is_on(rpr, "i")
        strike = _attr_is_on(rpr, "strike")

        rStyle = rpr.find(_w("rStyle"))
        if rStyle is not None and "Code" in rStyle.get(_w("val"), ""):
            is_code = True

        rFonts = rpr.find(_w("rFonts"))
        if rFonts is not None:
            for attr in (f"{{{_NS_W}}}ascii", f"{{{_NS_W}}}hAnsi"):
                fname = rFonts.get(attr, "")
                if "Courier" in fname or "Mono" in fname:
                    is_code = True

    if is_code:
        return f"`{text}`"
    if bold and italic:
        return f"***{text}***"
    if bold:
        return f"**{text}**"
    if italic:
        return f"*{text}*"
    if strike:
        return f"~~{text}~~"
    return text


def _para_to_md(para) -> str:
    """Convert a paragraph's runs and hyperlinks to inline Markdown."""
    rels: dict = {}
    try:
        for rel in para.part.rels.values():
            if "hyperlink" in rel.reltype:
                rels[rel.rId] = rel.target_ref
    except Exception:
        pass

    parts = []
    for child in para._p:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if local == "r":
            parts.append(_run_el_to_md(child))

        elif local == "hyperlink":
            r_id = child.get(f"{{{_NS_R}}}id")
            url = rels.get(r_id, "#") if r_id else "#"
            link_text = "".join(t.text or "" for t in child.iter(_w("t")))
            if link_text:
                parts.append(f"[{link_text}]({url})")

        elif local == "ins":
            # Tracked-change insertions — process inner runs
            for r_el in child.findall(_w("r")):
                parts.append(_run_el_to_md(r_el))

    return "".join(parts)


def _table_el_to_md(table) -> list[str]:
    """Convert a python-docx Table to a list of Markdown lines."""
    rows = table.rows
    if not rows:
        return []

    lines = []
    for row_idx, row in enumerate(rows):
        cells = [c.text.replace("\n", " ").replace("|", "\\|").strip()
                 for c in row.cells]
        lines.append("| " + " | ".join(cells) + " |")
        if row_idx == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return lines


def convert_docx_to_md(docx_path: Path, md_path: Path):
    """Convert a DOCX file to Markdown."""
    doc = Document(str(docx_path))

    output: list[str] = []
    code_buf: list[str] = []   # buffer for consecutive Code Block paragraphs

    def flush_code():
        if code_buf:
            output.append("```")
            output.extend(code_buf)
            output.append("```")
            output.append("")
            code_buf.clear()

    def add_blank():
        if output and output[-1] != "":
            output.append("")

    # Build a map from element to python-docx Table object
    table_map = {t._element: t for t in doc.tables}

    for child in doc.element.body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if local == "p":
            # Find matching paragraph object
            para = next((p for p in doc.paragraphs if p._element is child), None)
            if para is None:
                continue

            style = para.style.name
            raw_text = para.text

            # ── Code block paragraph ──────────────────────────────────────────
            if style == "Code Block":
                code_buf.append(raw_text)
                continue

            flush_code()

            if not raw_text.strip():
                add_blank()
                continue

            md_text = _para_to_md(para)

            # ── Headings ──────────────────────────────────────────────────────
            if re.match(r"Heading \d", style):
                level = int(style[-1])
                add_blank()
                output.append("#" * level + " " + md_text)
                output.append("")

            # ── Bullet list ───────────────────────────────────────────────────
            elif style.startswith("List Bullet"):
                suffix = style.replace("List Bullet", "").strip()
                level  = int(suffix) if suffix.isdigit() else 1
                output.append("  " * (level - 1) + "- " + md_text)

            # ── Numbered list ─────────────────────────────────────────────────
            elif style.startswith("List Number"):
                suffix = style.replace("List Number", "").strip()
                level  = int(suffix) if suffix.isdigit() else 1
                output.append("  " * (level - 1) + "1. " + md_text)

            # ── Blockquote ────────────────────────────────────────────────────
            elif style in ("Block Quote", "Quote"):
                output.append("> " + md_text)

            # ── Normal / everything else ──────────────────────────────────────
            else:
                output.append(md_text)

        elif local == "tbl":
            flush_code()
            table = table_map.get(child)
            if table is not None:
                add_blank()
                output.extend(_table_el_to_md(table))
                output.append("")

        elif local == "sectPr":
            pass  # section properties — ignore

    flush_code()

    # Remove trailing blank lines
    while output and output[-1] == "":
        output.pop()

    md_path.write_text("\n".join(output) + "\n", encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# PDF → MD conversion
# ─────────────────────────────────────────────────────────────────────────────

def _heading_level_from_size(size: float, body_size: float) -> int:
    """Return Markdown heading level (1-4) or 0 for body text."""
    if body_size <= 0:
        return 0
    ratio = size / body_size
    if ratio >= 1.8:
        return 1
    if ratio >= 1.4:
        return 2
    if ratio >= 1.2:
        return 3
    if ratio >= 1.1:
        return 4
    return 0


def _line_chars_to_text(chars: list) -> str:
    """Join a sorted list of char dicts into a string, inserting spaces at word gaps."""
    if not chars:
        return ""
    sorted_chars = sorted(chars, key=lambda c: c.get("x0", 0))
    result = []
    prev_x1 = None
    for ch in sorted_chars:
        text = ch.get("text", "")
        x0 = ch.get("x0", 0)
        size = float(ch.get("size", 10))
        if prev_x1 is not None and x0 - prev_x1 > size * 0.25:
            result.append(" ")
        result.append(text)
        prev_x1 = ch.get("x1", x0 + size * 0.5)
    return "".join(result).strip()


def _ocr_page_to_text(pdf_path: Path, page_number: int) -> str:
    """
    Run Tesseract OCR on a single PDF page (1-based index).
    Returns extracted plain text, or empty string on failure.
    Requires: pytesseract + tesseract binary + pdf2image + poppler.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path
        images = convert_from_path(
            str(pdf_path),
            dpi=300,
            first_page=page_number,
            last_page=page_number,
        )
        if not images:
            return ""
        text = pytesseract.image_to_string(images[0], lang="por+eng")
        return text.strip()
    except Exception:
        return ""


def _page_has_text(page) -> bool:
    """Return True if pdfplumber page contains meaningful extractable text."""
    chars = page.chars
    return len(chars) > 20  # threshold: at least 20 glyphs


def convert_pdf_to_md(pdf_path: Path, md_path: Path):
    """Convert a PDF file to Markdown using pdfplumber (with OCR fallback)."""
    try:
        import pdfplumber
    except ImportError:
        print("[ERRO] pdfplumber não instalado. Execute: pip install pdfplumber")
        sys.exit(1)

    all_output: list[str] = []

    with pdfplumber.open(str(pdf_path)) as pdf:

        # ── Pass 1: determine body font size (most frequent across all pages) ──
        size_freq: dict[float, int] = {}
        for page in pdf.pages:
            for ch in page.chars:
                s = round(float(ch.get("size", 12)), 1)
                size_freq[s] = size_freq.get(s, 0) + 1
        body_size = max(size_freq, key=size_freq.get) if size_freq else 12.0

        # ── Pass 2: page-by-page extraction ───────────────────────────────────
        for page_idx, page in enumerate(pdf.pages, start=1):
            page_out: list[str] = []

            # ── OCR fallback for scanned/image-only pages ─────────────────────
            if not _page_has_text(page):
                ocr_text = _ocr_page_to_text(pdf_path, page_idx)
                if ocr_text:
                    for line in ocr_text.splitlines():
                        line = line.strip()
                        if line:
                            page_out.append(line)
                        elif page_out and page_out[-1] != "":
                            page_out.append("")
                all_output.extend(page_out)
                all_output.append("")
                continue

            # Locate tables and pre-render them as Markdown
            try:
                table_objs = page.find_tables()
            except Exception:
                table_objs = []

            table_bboxes = [t.bbox for t in table_objs]  # (x0, top, x1, bottom)

            tables_by_y: dict[float, list[str]] = {}
            for t_obj in table_objs:
                data = t_obj.extract()
                if not data:
                    continue
                md_rows: list[str] = []
                for row_idx, row in enumerate(data):
                    cells = [
                        str(c or "").replace("\n", " ").replace("|", "\\|").strip()
                        for c in row
                    ]
                    md_rows.append("| " + " | ".join(cells) + " |")
                    if row_idx == 0:
                        md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
                tables_by_y[t_obj.bbox[1]] = md_rows

            # Filter chars that fall inside table bounding boxes
            def _in_table(ch) -> bool:
                cx0, ctop = ch.get("x0", 0), ch.get("top", 0)
                cx1, cbot = ch.get("x1", 0), ch.get("bottom", 0)
                for (tx0, ttop, tx1, tbot) in table_bboxes:
                    if cx0 >= tx0 - 2 and cx1 <= tx1 + 2 and ctop >= ttop - 2 and cbot <= tbot + 2:
                        return True
                return False

            text_chars = [ch for ch in page.chars if not _in_table(ch)]

            # Group chars into visual lines (tolerance ±2 pt on y)
            raw_lines: dict[float, list] = {}
            for ch in text_chars:
                top = ch.get("top", 0)
                matched = next((k for k in raw_lines if abs(k - top) <= 2), None)
                if matched is None:
                    raw_lines[top] = []
                    matched = top
                raw_lines[matched].append(ch)

            sorted_lines = sorted(raw_lines.items())  # ascending by y

            # Group consecutive lines into paragraphs:
            # same dominant font size + gap ≤ 1.6× font size → same paragraph
            paragraphs: list[list[tuple[float, list]]] = []
            for line_y, chars in sorted_lines:
                if not chars:
                    continue
                sizes = [float(ch.get("size", body_size)) for ch in chars]
                dom_size = max(set(sizes), key=sizes.count)

                merge = False
                if paragraphs:
                    last_line_y, last_chars = paragraphs[-1][-1]
                    if last_chars:
                        lsizes = [float(ch.get("size", body_size)) for ch in last_chars]
                        last_size = max(set(lsizes), key=lsizes.count)
                        gap = line_y - last_line_y
                        if abs(dom_size - last_size) < 1.5 and gap <= last_size * 1.6:
                            merge = True

                if merge:
                    paragraphs[-1].append((line_y, chars))
                else:
                    paragraphs.append([(line_y, chars)])

            # Render paragraphs, interleaving tables at their y positions
            table_ys = sorted(tables_by_y.keys())
            table_idx = 0

            for para_lines in paragraphs:
                para_y = para_lines[0][0]

                # Insert tables that appear above this paragraph
                while table_idx < len(table_ys) and table_ys[table_idx] < para_y:
                    page_out.append("")
                    page_out.extend(tables_by_y[table_ys[table_idx]])
                    page_out.append("")
                    table_idx += 1

                # Build paragraph text by joining each line
                line_texts = [_line_chars_to_text(chars) for _, chars in para_lines]
                text = " ".join(t for t in line_texts if t)
                if not text:
                    continue

                # Determine dominant attributes across the whole paragraph
                all_chars = [ch for _, chars in para_lines for ch in chars]
                sizes = [float(ch.get("size", body_size)) for ch in all_chars]
                dom_size = max(set(sizes), key=sizes.count)
                fontnames = [ch.get("fontname", "") for ch in all_chars]
                is_bold = any("Bold" in fn or "bold" in fn for fn in fontnames)
                is_italic = any(
                    "Italic" in fn or "Oblique" in fn or "italic" in fn
                    for fn in fontnames
                )

                heading = _heading_level_from_size(dom_size, body_size)

                if heading > 0:
                    page_out.append("")
                    page_out.append("#" * heading + " " + text)
                    page_out.append("")
                elif is_bold and not is_italic:
                    page_out.append(f"**{text}**")
                else:
                    if is_italic:
                        page_out.append(f"*{text}*")
                    else:
                        page_out.append(text)

            # Append remaining tables at page bottom
            while table_idx < len(table_ys):
                page_out.append("")
                page_out.extend(tables_by_y[table_ys[table_idx]])
                page_out.append("")
                table_idx += 1

            all_output.extend(page_out)
            all_output.append("")  # visual separator between pages

    # Collapse multiple consecutive blank lines into one
    cleaned: list[str] = []
    for line in all_output:
        if line == "" and cleaned and cleaned[-1] == "":
            continue
        cleaned.append(line)

    while cleaned and cleaned[-1] == "":
        cleaned.pop()

    md_path.write_text("\n".join(cleaned) + "\n", encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX footnote attachment (XML level — python-docx has no native support)
# ─────────────────────────────────────────────────────────────────────────────

_W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_NS = "http://www.w3.org/XML/1998/namespace"

_FOOTNOTES_CT  = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
_FOOTNOTES_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
_FOOTNOTES_URI = "/word/footnotes.xml"


def _fn_el(tag: str, parent=None, attribs: dict | None = None):
    """Create a w:-namespaced element, optionally appending it to parent."""
    el = OxmlElement(f"w:{tag}")
    if attribs:
        for k, v in attribs.items():
            el.set(qn(f"w:{k}"), v)
    if parent is not None:
        parent.append(el)
    return el


def _render_inline_tokens_to_runs(renderer: "DocxRenderer", tokens: list) -> list:
    """
    Render inline tokens via a throw-away paragraph and return deep-copied XML elements.
    Hyperlinks are flattened to plain runs to avoid cross-part relationship issues.
    """
    import copy
    from docx import Document as _TmpDoc

    tmp_doc  = _TmpDoc()
    tmp_para = tmp_doc.add_paragraph()
    renderer._render_inline_to_para(tmp_para, tokens)

    runs = []
    for child in tmp_para._p:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "r":
            runs.append(copy.deepcopy(child))
        elif local == "hyperlink":
            # Flatten hyperlink runs to avoid cross-part rels
            for r_el in child.findall(_w("r")):
                runs.append(copy.deepcopy(r_el))
    return runs


def _build_footnotes_element(footnotes: list[tuple[int, list]]):
    """Build the <w:footnotes> lxml element tree with rich run content."""
    from lxml import etree

    root = etree.Element(
        f"{{{_W_NS}}}footnotes",
        nsmap={
            "w":   _W_NS,
            "xml": _XML_NS,
        },
    )

    # Required separator stubs (id -1 and 0)
    for fn_id, fn_type, sep_tag in [
        ("-1", "separator",             "separator"),
        ("0",  "continuationSeparator", "continuationSeparator"),
    ]:
        fn_el  = etree.SubElement(root, f"{{{_W_NS}}}footnote")
        fn_el.set(f"{{{_W_NS}}}type", fn_type)
        fn_el.set(f"{{{_W_NS}}}id",   fn_id)
        p      = etree.SubElement(fn_el, f"{{{_W_NS}}}p")
        pPr    = etree.SubElement(p, f"{{{_W_NS}}}pPr")
        sp     = etree.SubElement(pPr, f"{{{_W_NS}}}spacing")
        sp.set(f"{{{_W_NS}}}after",    "0")
        sp.set(f"{{{_W_NS}}}line",     "240")
        sp.set(f"{{{_W_NS}}}lineRule", "auto")
        r      = etree.SubElement(p, f"{{{_W_NS}}}r")
        etree.SubElement(r, f"{{{_W_NS}}}{sep_tag}")

    # Actual footnotes — fn_runs is a list of pre-rendered lxml <w:r> elements
    for fn_id, fn_runs in footnotes:
        fn_el  = etree.SubElement(root, f"{{{_W_NS}}}footnote")
        fn_el.set(f"{{{_W_NS}}}id", str(fn_id))
        p      = etree.SubElement(fn_el, f"{{{_W_NS}}}p")
        pPr    = etree.SubElement(p, f"{{{_W_NS}}}pPr")
        pStyle = etree.SubElement(pPr, f"{{{_W_NS}}}pStyle")
        pStyle.set(f"{{{_W_NS}}}val", "FootnoteText")

        # Footnote mark (superscript anchor inside footnote paragraph)
        r_mark   = etree.SubElement(p, f"{{{_W_NS}}}r")
        rPr_m    = etree.SubElement(r_mark, f"{{{_W_NS}}}rPr")
        rStyle_m = etree.SubElement(rPr_m, f"{{{_W_NS}}}rStyle")
        rStyle_m.set(f"{{{_W_NS}}}val", "FootnoteReference")
        etree.SubElement(r_mark, f"{{{_W_NS}}}footnoteRef")

        # Leading space run (conventional)
        r_sp = etree.SubElement(p, f"{{{_W_NS}}}r")
        t_sp = etree.SubElement(r_sp, f"{{{_W_NS}}}t")
        t_sp.set(f"{{{_XML_NS}}}space", "preserve")
        t_sp.text = " "

        # Rich content runs
        for run_el in fn_runs:
            p.append(run_el)

    return root


def _attach_footnotes(doc: Document, footnotes_raw: list[tuple[int, list]],
                      renderer: "DocxRenderer"):
    """Pre-render inline tokens, build footnotes.xml, and relate it to the document."""
    if not footnotes_raw:
        return

    from docx.opc.part import XmlPart
    from docx.opc.packuri import PackURI

    # Convert inline tokens → pre-rendered XML run elements
    footnotes_rendered = [
        (fn_id, _render_inline_tokens_to_runs(renderer, tokens))
        for fn_id, tokens in footnotes_raw
    ]

    fn_element = _build_footnotes_element(footnotes_rendered)
    part_uri   = PackURI(_FOOTNOTES_URI)
    fn_part    = XmlPart(part_uri, _FOOTNOTES_CT, fn_element, doc.part.package)
    doc.part.relate_to(fn_part, _FOOTNOTES_REL)


# ─────────────────────────────────────────────────────────────────────────────
# Core conversion function
# ─────────────────────────────────────────────────────────────────────────────

def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Convert a single Markdown file to DOCX."""
    md_text = md_path.read_text(encoding="utf-8", errors="replace")

    # Strip YAML front matter
    md_text = re.sub(r"^---\s*\n.*?\n---\s*\n", "", md_text, flags=re.DOTALL)

    # Extract cover block (title, subtitle, author) before main content
    cover, md_body = extract_cover(md_text)

    doc = Document()

    # Page setup — A4
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    setup_styles(doc)
    add_footer_page_numbers(doc)

    # Remove the leading empty paragraph Word always creates
    if doc.paragraphs and doc.paragraphs[0].text == "":
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Cover page (if found)
    if cover:
        add_cover_page(doc, cover)
        add_toc(doc)

    renderer = DocxRenderer(doc=doc, md_path=md_path)
    md_parser = mistune.create_markdown(
        renderer=renderer,
        plugins=["strikethrough", "table", "footnotes", "task_lists"],
    )
    md_parser(md_body)

    if renderer._footnotes:
        _attach_footnotes(doc, renderer._footnotes, renderer)

    doc.save(str(docx_path))


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────

def main():
    print("MD ↔ DOCX Converter by Jair Lima")
    print("=" * 34)

    parser = argparse.ArgumentParser(
        description="MD ↔ DOCX / PDF → MD Converter by Jair Lima",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  md2docx                           # Convert all .md in current folder
  md2docx README.md                 # Convert .md → .docx
  md2docx relatorio.docx            # Convert .docx → .md
  md2docx artigo.pdf                # Convert .pdf  → .md
  md2docx README.md --force         # Overwrite existing file
  md2docx --folder /path/to/dir     # Convert all .md in a folder
  md2docx /path/to/dir              # Same — directory as positional arg
  md2docx --folder . --recursive    # Include subfolders
        """,
    )
    parser.add_argument(
        "file",
        nargs="?",
        help="Specific .md file to convert, or a folder path",
    )
    parser.add_argument(
        "--folder",
        default=".",
        help="Folder to scan for .md files (default: current directory)",
    )
    parser.add_argument(
        "--recursive", "-r",
        action="store_true",
        help="Scan subfolders recursively for .md files",
    )
    parser.add_argument(
        "--force", "-f",
        action="store_true",
        help="Overwrite existing DOCX files",
    )
    parser.add_argument(
        "--output", "-o",
        help="Output folder for DOCX files (default: same as source)",
    )
    parser.add_argument(
        "--version", "-v",
        action="version",
        version=f"MD to DOCX Converter by Jair Lima v{VERSION}",
    )
    args = parser.parse_args()

    output_dir = Path(args.output) if args.output else None

    # ── Single file mode ──────────────────────────────────────────────────────
    if args.file:
        src_path = Path(args.file).resolve()

        # If argument is a directory, redirect to folder scan mode
        if src_path.is_dir():
            args.folder = str(src_path)
        else:
            if not src_path.exists():
                print(f"[ERRO] Arquivo não encontrado: {src_path}")
                sys.exit(1)

            ext = src_path.suffix.lower()

            # ── PDF → MD ──────────────────────────────────────────────────────
            if ext == ".pdf":
                out_dir = output_dir or src_path.parent
                out_dir.mkdir(parents=True, exist_ok=True)
                md_out = out_dir / (src_path.stem + ".md")

                if md_out.exists() and not args.force:
                    print(f"[PULA]  {src_path.name} → já existe {md_out.name}")
                    sys.exit(0)

                print(f"[CONV]  {src_path.name} → {md_out.name}  (PDF → MD)")
                with Spinner(f"{src_path.name}"):
                    t0 = time.time()
                    convert_pdf_to_md(src_path, md_out)
                    elapsed = time.time() - t0
                size_kb = md_out.stat().st_size / 1024
                print(f"[OK]    Salvo em: {md_out}  ({elapsed:.1f}s, {size_kb:.0f} KB)")
                sys.exit(0)

            # ── DOCX/DOC → MD (modo inverso) ──────────────────────────────────
            if ext in (".docx", ".doc"):
                if ext == ".doc":
                    print(f"[ERRO] Arquivos .doc legados não são suportados. Abra no Word e salve como .docx.")
                    sys.exit(1)

                out_dir = output_dir or src_path.parent
                out_dir.mkdir(parents=True, exist_ok=True)
                md_out = out_dir / (src_path.stem + ".md")

                if md_out.exists() and not args.force:
                    print(f"[PULA]  {src_path.name} → já existe {md_out.name}")
                    sys.exit(0)

                print(f"[CONV]  {src_path.name} → {md_out.name}  (DOCX → MD)")
                with Spinner(f"{src_path.name}"):
                    t0 = time.time()
                    convert_docx_to_md(src_path, md_out)
                    elapsed = time.time() - t0
                size_kb = md_out.stat().st_size / 1024
                print(f"[OK]    Salvo em: {md_out}  ({elapsed:.1f}s, {size_kb:.0f} KB)")
                sys.exit(0)

            # ── MD → DOCX (modo normal) ────────────────────────────────────────
            if ext != ".md":
                print(f"[ERRO] Extensão não suportada: '{ext}'. Use .md, .docx, .pdf ou uma pasta.")
                sys.exit(1)

            md_path = src_path
            out_dir = output_dir or md_path.parent
            out_dir.mkdir(parents=True, exist_ok=True)
            docx_path = out_dir / (md_path.stem + ".docx")

            if docx_path.exists() and not args.force:
                print(f"[PULA]  {md_path.name} → já existe {docx_path.name}")
                sys.exit(0)

            print(f"[CONV]  {md_path.name} → {docx_path.name}")
            with Spinner(f"{md_path.name}"):
                t0 = time.time()
                convert_md_to_docx(md_path, docx_path)
                elapsed = time.time() - t0
            size_kb = docx_path.stat().st_size / 1024
            print(f"[OK]    Salvo em: {docx_path}  ({elapsed:.1f}s, {size_kb:.0f} KB)")
            sys.exit(0)

    # ── Folder scan mode ──────────────────────────────────────────────────────
    folder = Path(args.folder).resolve()
    if not folder.is_dir():
        print(f"[ERRO] Pasta não encontrada: {folder}")
        sys.exit(1)

    if args.recursive:
        md_files = sorted(folder.rglob("*.md"))
    else:
        md_files = sorted(folder.glob("*.md"))

    if not md_files:
        suffix = " (e subpastas)" if args.recursive else ""
        print(f"[INFO] Nenhum arquivo .md encontrado em: {folder}{suffix}")
        sys.exit(0)

    converted = 0
    skipped = 0
    errors = 0

    rec_label = " (recursivo)" if args.recursive else ""
    print(f"[SCAN]  {folder}{rec_label}")
    print(f"[INFO]  {len(md_files)} arquivo(s) .md encontrado(s)\n")

    for md_path in md_files:
        # Preserve subfolder structure when --output is set
        if output_dir:
            rel = md_path.relative_to(folder)
            out_dir = output_dir / rel.parent
        else:
            out_dir = md_path.parent
        out_dir.mkdir(parents=True, exist_ok=True)
        docx_path = out_dir / (md_path.stem + ".docx")

        display_name = str(md_path.relative_to(folder)) if args.recursive else md_path.name

        if docx_path.exists() and not args.force:
            print(f"  [PULA]  {display_name}")
            skipped += 1
            continue

        try:
            print(f"  [CONV]  {display_name} → {docx_path.name}")
            with Spinner(md_path.name):
                t0 = time.time()
                convert_md_to_docx(md_path, docx_path)
                elapsed = time.time() - t0
            size_kb = docx_path.stat().st_size / 1024
            print(f"  [OK]    Salvo. ({elapsed:.1f}s, {size_kb:.0f} KB)")
            converted += 1
        except Exception as e:
            print(f"  [ERRO]  {md_path.name}: {e}")
            errors += 1

    print(f"\n[FIM]   Convertidos: {converted}  |  Pulados: {skipped}  |  Erros: {errors}")


if __name__ == "__main__":
    main()
