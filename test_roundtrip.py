"""
Teste de regressão do md2docx by Jair Lima.

Converte fixtures MD -> DOCX (e DOCX -> MD de volta) e verifica que a
formatação sobrevive. Criado depois de 3 bugs silenciosos na v3.4/v3.5:
itálico/tachado descartados, cabeçalho de tabela duplicado e numeração
de lista compartilhada entre listas diferentes.

Uso: python test_roundtrip.py  (exit code 0 = tudo passou, 1 = falhou)
"""
import sys
import tempfile
from pathlib import Path

import docx

import md2docx as m

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

FAILURES: list[str] = []


def check(condition: bool, description: str) -> None:
    status = "OK" if condition else "FALHOU"
    print(f"  [{status}] {description}")
    if not condition:
        FAILURES.append(description)


def numid_of(paragraph) -> str | None:
    numid_el = paragraph._p.find(f".//{W_NS}numId")
    return numid_el.get(f"{W_NS}val") if numid_el is not None else None


def test_inline_formatting(tmp_dir: Path) -> None:
    print("\n== Formatação inline (negrito, itálico, negrito+itálico, tachado, código) ==")
    md_path = tmp_dir / "inline.md"
    md_path.write_text(
        "Texto com **negrito**, *itálico*, ***negrito e itálico***, "
        "~~tachado~~ e `código inline` no meio.\n",
        encoding="utf-8",
    )
    docx_path = tmp_dir / "inline.docx"
    m.convert_md_to_docx(md_path, docx_path)

    doc = docx.Document(str(docx_path))
    body_para = next(p for p in doc.paragraphs if "Texto com" in p.text)
    runs = body_para.runs

    bold_run = next((r for r in runs if r.text == "negrito" and r.bold and not r.italic), None)
    italic_run = next((r for r in runs if r.text == "itálico" and r.italic and not r.bold), None)
    bold_italic_run = next((r for r in runs if r.text == "negrito e itálico" and r.bold and r.italic), None)
    strike_run = next((r for r in runs if r.text == "tachado" and r.font.strike), None)
    code_run = next((r for r in runs if r.text == "código inline"), None)

    check(bold_run is not None, "negrito preservado como run bold")
    check(italic_run is not None, "itálico preservado como run italic (mistune 'emphasis')")
    check(bold_italic_run is not None, "negrito+itálico combinados preservados")
    check(strike_run is not None, "tachado preservado como run strike (mistune 'strikethrough')")
    check(code_run is not None and code_run.font.name == "Courier New", "código inline preservado com fonte monoespaçada")

    # round-trip de volta para MD
    md_back_path = tmp_dir / "inline_back.md"
    m.convert_docx_to_md(docx_path, md_back_path)
    md_back = md_back_path.read_text(encoding="utf-8")
    check("**negrito**" in md_back, "round-trip: **negrito** volta ao MD")
    check("*itálico*" in md_back and "***" not in md_back.split("*itálico*")[0][-3:], "round-trip: *itálico* volta ao MD")
    check("***negrito e itálico***" in md_back, "round-trip: ***negrito e itálico*** volta ao MD")
    check("~~tachado~~" in md_back, "round-trip: ~~tachado~~ volta ao MD")
    check("`código inline`" in md_back, "round-trip: `código inline` volta ao MD")


def test_table_header_not_duplicated(tmp_dir: Path) -> None:
    print("\n== Tabela (cabeçalho não duplicado) ==")
    md_path = tmp_dir / "tabela.md"
    md_path.write_text(
        "| Coluna A | Coluna B |\n"
        "|----------|----------|\n"
        "| a1       | b1       |\n"
        "| a2       | b2       |\n",
        encoding="utf-8",
    )
    docx_path = tmp_dir / "tabela.docx"
    m.convert_md_to_docx(md_path, docx_path)

    doc = docx.Document(str(docx_path))
    table = doc.tables[0]
    check(len(table.rows) == 3, f"tabela tem 3 linhas (1 cabeçalho + 2 dados), tem {len(table.rows)}")
    check(table.rows[0].cells[0].text.strip() == "Coluna A", "cabeçalho aparece só uma vez, sem linha vazia duplicada")


def test_list_numbering_restarts(tmp_dir: Path) -> None:
    print("\n== Numeração de listas (cada lista reinicia em 1) ==")
    md_path = tmp_dir / "listas.md"
    md_path.write_text(
        "# Capitulo 1\n\nTexto.\n\n1. Item A\n2. Item B\n3. Item C\n\n"
        "# Capitulo 2\n\nTexto.\n\n1. Item X\n2. Item Y\n\n"
        "# Capitulo 3\n\n1. Item Alfa\n2. Item Beta\n3. Item Gama\n4. Item Delta\n",
        encoding="utf-8",
    )
    docx_path = tmp_dir / "listas.docx"
    m.convert_md_to_docx(md_path, docx_path)

    doc = docx.Document(str(docx_path))
    list_paras = [p for p in doc.paragraphs if p.style.name.startswith("List Number")]
    check(len(list_paras) == 9, f"9 itens de lista numerada no total, encontrados {len(list_paras)}")

    numids = [numid_of(p) for p in list_paras]
    check(all(n is not None for n in numids), "todo item de lista numerada tem numId explícito (não depende só do estilo)")

    distinct_numids = set(numids)
    check(len(distinct_numids) == 3, f"3 listas distintas usam 3 numIds diferentes, encontrados {len(distinct_numids)}")

    # cada lista mantém o MESMO numId internamente (A,B,C compartilham; X,Y compartilham; Alfa..Delta compartilham)
    check(numids[0] == numids[1] == numids[2], "lista do Capitulo 1 usa um único numId consistente")
    check(numids[3] == numids[4], "lista do Capitulo 2 usa um único numId consistente")
    check(numids[5] == numids[6] == numids[7] == numids[8], "lista do Capitulo 3 usa um único numId consistente")
    check(numids[0] != numids[3] != numids[5] and numids[0] != numids[5], "as 3 listas não compartilham numId entre si")


def test_nested_ordered_list(tmp_dir: Path) -> None:
    print("\n== Lista numerada aninhada (sublista independente) ==")
    md_path = tmp_dir / "aninhado.md"
    md_path.write_text(
        "1. Primeiro passo\n2. Segundo passo\n   1. Sub-passo 2.1\n   2. Sub-passo 2.2\n3. Terceiro passo\n",
        encoding="utf-8",
    )
    docx_path = tmp_dir / "aninhado.docx"
    m.convert_md_to_docx(md_path, docx_path)

    doc = docx.Document(str(docx_path))
    top_level = [p for p in doc.paragraphs if p.style.name == "List Number"]
    nested = [p for p in doc.paragraphs if p.style.name == "List Number 2"]

    check(len(top_level) == 3, f"3 itens no nível 1 (Primeiro/Segundo/Terceiro), encontrados {len(top_level)}")
    check(len(nested) == 2, f"2 itens no nível 2 (Sub-passo 2.1/2.2), encontrados {len(nested)}")

    top_numids = {numid_of(p) for p in top_level}
    nested_numids = {numid_of(p) for p in nested}
    check(len(top_numids) == 1, "nível 1 mantém um único numId (não reinicia no meio por causa da sublista)")
    check(len(nested_numids) == 1, "nível 2 (sublista) tem seu próprio numId único")
    check(top_numids != nested_numids, "sublista usa numId diferente do nível pai")


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="md2docx_test_") as tmp:
        tmp_dir = Path(tmp)
        test_inline_formatting(tmp_dir)
        test_table_header_not_duplicated(tmp_dir)
        test_list_numbering_restarts(tmp_dir)
        test_nested_ordered_list(tmp_dir)

    print()
    if FAILURES:
        print(f"{len(FAILURES)} verificação(ões) falharam:")
        for f in FAILURES:
            print(f"  - {f}")
        return 1

    print("Todas as verificações passaram.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
